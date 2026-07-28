import io
from docx import Document
from src.parser import parse_document

def test_txt_parse():
    result = parse_document(b"Hello resume", "resume.txt")
    assert result.text == "Hello resume"

def test_docx_parse_includes_tables():
    buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Education")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Python"
    doc.save(buffer)
    result = parse_document(buffer.getvalue(), "resume.docx")
    assert "Python" in result.text
