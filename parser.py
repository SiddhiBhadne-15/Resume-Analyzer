"""Safe text extraction from uploaded resume and job-description files."""
from __future__ import annotations

import io
import os
import re
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO

from docx import Document
from PyPDF2 import PdfReader


@dataclass
class ParseResult:
    text: str
    file_type: str
    pages: int | None = None
    warnings: list[str] | None = None


class DocumentParseError(ValueError):
    pass


def _clean_extracted_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_bytes(file: bytes | BinaryIO) -> bytes:
    if isinstance(file, bytes):
        return file
    file.seek(0)
    return file.read()


def parse_pdf(data: bytes) -> ParseResult:
    warnings: list[str] = []
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise DocumentParseError("The PDF is password-protected.") from exc
        page_text = []
        for index, page in enumerate(reader.pages):
            try:
                page_text.append(page.extract_text() or "")
            except Exception:
                warnings.append(f"Could not extract text from page {index + 1}.")
        text = _clean_extracted_text("\n".join(page_text))
        if not text:
            raise DocumentParseError(
                "No selectable text was found. The PDF may be scanned; run OCR and upload it again."
            )
        return ParseResult(text, "pdf", len(reader.pages), warnings)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Unable to read PDF: {exc}") from exc


def parse_docx(data: bytes) -> ParseResult:
    try:
        doc = Document(io.BytesIO(data))
        chunks = [p.text for p in doc.paragraphs]
        # Text inside tables is common in resumes.
        for table in doc.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text for cell in row.cells))
        text = _clean_extracted_text("\n".join(chunks))
        if not text:
            raise DocumentParseError("No text was found in the DOCX file.")
        return ParseResult(text, "docx", None, [])
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Unable to read DOCX: {exc}") from exc


def parse_legacy_doc(data: bytes) -> ParseResult:
    """Parse .doc through the system's antiword command when available."""
    path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp:
            temp.write(data)
            path = temp.name
        result = subprocess.run(
            ["antiword", path], capture_output=True, text=True, timeout=20, check=False
        )
        if result.returncode != 0 or not result.stdout.strip():
            raise DocumentParseError(
                "Legacy .doc extraction needs the 'antiword' system package. "
                "For best results, save the file as DOCX or PDF."
            )
        return ParseResult(_clean_extracted_text(result.stdout), "doc", None, [])
    except FileNotFoundError as exc:
        raise DocumentParseError(
            "Legacy .doc extraction is unavailable. Save the document as DOCX or PDF."
        ) from exc
    finally:
        if path and os.path.exists(path):
            os.unlink(path)


def parse_document(file: bytes | BinaryIO, filename: str) -> ParseResult:
    data = _read_bytes(file)
    if not data:
        raise DocumentParseError("The uploaded file is empty.")
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(data)
    if suffix == ".docx":
        return parse_docx(data)
    if suffix == ".doc":
        return parse_legacy_doc(data)
    if suffix in {".txt", ".md"}:
        try:
            text = _clean_extracted_text(data.decode("utf-8-sig"))
        except UnicodeDecodeError:
            text = _clean_extracted_text(data.decode("latin-1"))
        if not text:
            raise DocumentParseError("The uploaded text file is empty.")
        return ParseResult(text, suffix.lstrip("."), None, [])
    raise DocumentParseError("Unsupported format. Upload PDF, DOCX, DOC, or TXT.")
